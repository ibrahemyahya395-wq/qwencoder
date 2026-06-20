#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
PDF to Word Converter with OCR Support
Converts PDF files to Word documents with high accuracy.
Supports text-only and image+text PDFs.
"""

import os
import sys
import tempfile
from pathlib import Path
from datetime import datetime

from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def extract_text_from_image(image: Image) -> str:
    """Extract text from an image using Tesseract OCR."""
    # Use high-resolution settings for better accuracy
    custom_config = r'--oem 3 --psm 6 -l eng+ara'
    text = pytesseract.image_to_string(image, config=custom_config)
    return text.strip()


def process_pdf_page(page_image: Image, page_num: int) -> tuple[str, bool]:
    """
    Process a single PDF page.
    Returns: (extracted_text, has_images)
    """
    # First try to extract text directly if available
    text = extract_text_from_image(page_image)
    
    # Check if the page contains significant visual content
    # Convert to grayscale and check variance to detect images
    gray = page_image.convert('L')
    has_visual_content = True  # Assume all pages have visual content for now
    
    return text, has_visual_content


def convert_pdf_to_word(pdf_path: str, output_path: str = None) -> str:
    """
    Convert PDF to Word document.
    
    Args:
        pdf_path: Path to input PDF file
        output_path: Optional path for output Word file
        
    Returns:
        Path to generated Word document
    """
    pdf_path = Path(pdf_path)
    
    if not pdf_path.exists():
        raise FileNotFoundError(f"PDF file not found: {pdf_path}")
    
    # Generate output path if not provided
    if output_path is None:
        output_path = pdf_path.with_suffix('.docx')
    else:
        output_path = Path(output_path)
    
    print(f"Processing PDF: {pdf_path.name}")
    print(f"Output will be saved to: {output_path.name}")
    
    # Create Word document
    doc = Document()
    
    # Add title
    title = doc.add_heading(pdf_path.stem, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Converted on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Source file: {pdf_path.name}")
    doc.add_page_break()
    
    # Convert PDF pages to images
    print("Converting PDF pages to images...")
    try:
        # Convert with high DPI for better quality
        images = convert_from_path(pdf_path, dpi=300)
        print(f"Found {len(images)} pages")
    except Exception as e:
        raise RuntimeError(f"Failed to convert PDF to images: {str(e)}")
    
    # Process each page
    total_pages = len(images)
    for i, page_image in enumerate(images, 1):
        print(f"Processing page {i}/{total_pages}...")
        
        # Add page header
        doc.add_heading(f"Page {i}", level=1)
        
        # Extract text from the page
        text, has_visual = process_pdf_page(page_image, i)
        
        if text:
            # Add extracted text
            doc.add_paragraph(text)
        
        # Save the page image temporarily and add to document
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp_file:
            temp_path = tmp_file.name
            page_image.save(temp_path, 'PNG')
            
            # Add image to document
            try:
                # Resize image to fit page width
                img_width = 6.5  # inches (fits within standard page margins)
                doc.add_picture(temp_path, width=Inches(img_width))
                doc.add_paragraph().paragraph_format.space_after = Pt(12)
            finally:
                # Clean up temporary file
                os.unlink(temp_path)
        
        # Add page break between pages (except after last page)
        if i < total_pages:
            doc.add_page_break()
    
    # Save the document
    doc.save(output_path)
    print(f"\n✓ Successfully converted '{pdf_path.name}' to '{output_path.name}'")
    
    return str(output_path)


def main():
    """Main entry point for command-line usage."""
    if len(sys.argv) < 2:
        print("Usage:")
        print(f"  python {sys.argv[0]} <input.pdf> [output.docx]")
        print("\nExample:")
        print(f"  python {sys.argv[0]} document.pdf")
        print(f"  python {sys.argv[0]} document.pdf output.docx")
        sys.exit(1)
    
    input_pdf = sys.argv[1]
    output_docx = sys.argv[2] if len(sys.argv) > 2 else None
    
    try:
        result_path = convert_pdf_to_word(input_pdf, output_docx)
        print(f"\nOutput file: {result_path}")
    except Exception as e:
        print(f"Error: {str(e)}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    main()
